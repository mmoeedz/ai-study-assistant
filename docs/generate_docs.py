"""
Generates two professional Word documents for the AI Study Assistant project:

  1. AI_Study_Assistant_Proposal.docx
  2. AI_Study_Assistant_Guide.docx

Run:  python docs/generate_docs.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).resolve().parent.parent

# Brand palette (matches the app's UI)
GOLD = RGBColor(0xD4, 0xA8, 0x4B)
NAVY = RGBColor(0x0F, 0x1A, 0x2E)
INK = RGBColor(0x2B, 0x24, 0x18)
SOFT = RGBColor(0x55, 0x60, 0x6B)

# ── Modern proposal palette (matches the reference template) ────────
BLACK = RGBColor(0x12, 0x12, 0x12)
TEAL = RGBColor(0x1E, 0x7A, 0x8C)        # accent rule
TEAL_HEX = "1E7A8C"
LIGHT_BLUE_HEX = "7DB7C7"
DARK_HEX = "121212"
GRAY = RGBColor(0x6B, 0x72, 0x80)


# ─── helpers ─────────────────────────────────────────────────────────
def shade_cell(cell, color_hex: str) -> None:
    """Set cell background color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex: str = "D4A84B") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tc_pr.append(borders)


def add_heading(doc, text: str, level: int = 1, color=NAVY) -> None:
    h = doc.add_heading(level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(13)


def add_para(doc, text: str, *, size: int = 11, bold: bool = False, italic: bool = False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text: str, *, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = ""
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = "Calibri"


def add_numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text).font.size = Pt(11)


def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run("──────────────────────────────────────────────")
    run.font.color.rgb = GOLD
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_callout(doc, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, "FFF6DC")
    set_cell_borders(cell, "D4A84B")
    para = cell.paragraphs[0]
    t_run = para.add_run(f"{title}\n")
    t_run.bold = True
    t_run.font.size = Pt(12)
    t_run.font.color.rgb = NAVY
    b_run = para.add_run(body)
    b_run.font.size = Pt(11)
    b_run.font.color.rgb = INK
    doc.add_paragraph()  # spacing


def add_kv_table(doc, rows, headers=("Item", "Description")):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 5"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "0F1A2E")
    for row in rows:
        tr = table.add_row().cells
        for i, val in enumerate(row):
            tr[i].text = ""
            run = tr[i].paragraphs[0].add_run(val)
            run.font.size = Pt(10.5)
    doc.add_paragraph()


def setup_document() -> Document:
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def _center_run(doc, text, *, size, color=None, bold=False, italic=False, font="Calibri"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def _info_table(doc, rows):
    """Two-column borderless info table for title-page metadata."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = 1  # center
    for i, (k, v) in enumerate(rows):
        kc, vc = table.rows[i].cells
        kc.text = ""
        vc.text = ""
        # left column (label)
        rk = kc.paragraphs[0].add_run(k.upper())
        rk.font.name = "Calibri"
        rk.font.size = Pt(10.5)
        rk.font.color.rgb = SOFT
        rk.bold = True
        kc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # right column (value)
        rv = vc.paragraphs[0].add_run(v)
        rv.font.name = "Calibri"
        rv.font.size = Pt(11.5)
        rv.font.color.rgb = NAVY
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Make borders invisible
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                borders.append(b)
            tc_pr.append(borders)
    doc.add_paragraph()


def add_title_page(
    doc,
    title: str,
    subtitle: str,
    *,
    university: str = "Punjab University College of Information Technology (PUCIT)",
    course: str = "Natural Language Processing (NLP)",
    instructor: str = "Sir Adnan Abid",
    submission_label: str = "Project Proposal",
    team: list[tuple[str, str]] | None = None,
):
    """Render a professional academic cover page."""
    # ── Top university band ──
    _center_run(doc, university.upper(), size=12, color=GOLD, bold=True, font="Calibri")
    _center_run(doc, "Department of Data Science", size=10.5, color=SOFT, italic=True)

    # gold rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 28)
    run.font.color.rgb = GOLD
    run.font.size = Pt(11)

    # ── Crest spacer ──
    doc.add_paragraph()
    _center_run(doc, "🎓", size=46)

    # ── Title block ──
    _center_run(doc, title, size=34, color=NAVY, bold=True, font="Calibri")
    _center_run(doc, subtitle, size=13, color=GOLD, italic=True)

    # decorative star
    doc.add_paragraph()
    _center_run(doc, "─── ✦ ───", size=14, color=GOLD)
    doc.add_paragraph()

    # ── Submission label ──
    _center_run(doc, submission_label.upper(), size=14, color=NAVY, bold=True)
    doc.add_paragraph()

    # ── Course / instructor / submitted-on info ──
    _info_table(doc, [
        ("Course", course),
        ("Instructor", instructor),
        ("Institution", university),
    ])

    # ── Team table ──
    if team:
        doc.add_paragraph()
        _center_run(doc, "Submitted by", size=12, color=SOFT, italic=True)
        team_table = doc.add_table(rows=len(team), cols=2)
        team_table.alignment = 1
        for i, (rid, name) in enumerate(team):
            ic, nc = team_table.rows[i].cells
            ic.text = ""
            nc.text = ""
            r1 = ic.paragraphs[0].add_run(rid)
            r1.font.size = Pt(11.5)
            r1.font.color.rgb = GOLD
            r1.bold = True
            ic.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r2 = nc.paragraphs[0].add_run(name)
            r2.font.size = Pt(12)
            r2.font.color.rgb = NAVY
            r2.bold = True
            nc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # No borders
        for row in team_table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right"):
                    b = OxmlElement(f"w:{side}")
                    b.set(qn("w:val"), "nil")
                    borders.append(b)
                tc_pr.append(borders)

    # ── Footer rule ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 28)
    run.font.color.rgb = GOLD
    run.font.size = Pt(11)
    _center_run(doc, "BS Data Science  •  2023 Batch", size=10, color=SOFT, italic=True)

    doc.add_page_break()


# ─── Modern proposal helpers (matches the reference template) ───────
def _set_paragraph_border(paragraph, *, color=TEAL_HEX, size_pt=18, position="bottom"):
    """Add a coloured border below (or top of) a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bdr = OxmlElement(f"w:{position}")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), str(size_pt))   # eighths of a point
    bdr.set(qn("w:space"), "1")
    bdr.set(qn("w:color"), color)
    pBdr.append(bdr)


def add_modern_cover(
    doc,
    *,
    big_title_lines: tuple[str, str],
    author_lines: list[str],
    project_title: str,
    description: str,
    objective: str,
    opportunity: str,
    course_info: list[tuple[str, str]],
    team: list[tuple[str, str]],
):
    """Render a single-page modern proposal cover matching the reference."""

    # ── Top header: 2-col table (big stacked title | meta block) ────
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(11)
    table.columns[1].width = Cm(6)
    left, right = table.rows[0].cells
    left.width = Cm(11)
    right.width = Cm(6)

    # left: stacked PROJECT / PROPOSAL
    left.text = ""
    for line in big_title_lines:
        p = left.add_paragraph()
        run = p.add_run(line.upper())
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(38)
        run.font.color.rgb = BLACK
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 0.95

    # right: author block, top-aligned
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right.text = ""
    spacer = right.paragraphs[0]
    spacer.add_run("")  # vertical spacer
    spacer.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(author_lines):
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(11.5)
        run.font.color.rgb = BLACK
        run.bold = (i < len(author_lines) - 1)  # last line not bold
        p.paragraph_format.space_after = Pt(2)

    # remove all borders from header table
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                borders.append(b)
            tc_pr.append(borders)

    # ── Teal underline rule ────────────────────────────────────────
    rule_p = doc.add_paragraph()
    _set_paragraph_border(rule_p, color=TEAL_HEX, size_pt=18, position="bottom")
    rule_p.paragraph_format.space_before = Pt(2)
    rule_p.paragraph_format.space_after = Pt(20)

    # ── PROJECT TITLE ──────────────────────────────────────────────
    add_modern_section_heading(doc, "Project Title")
    add_para(doc, project_title, size=12)

    # ── DESCRIPTION ────────────────────────────────────────────────
    add_modern_section_heading(doc, "Description")
    add_para(doc, description, size=11)

    # ── OBJECTIVE ──────────────────────────────────────────────────
    add_modern_section_heading(doc, "Objective")
    add_para(doc, objective, size=11)

    # ── OPPORTUNITY ────────────────────────────────────────────────
    add_modern_section_heading(doc, "Opportunity")
    add_para(doc, opportunity, size=11)

    # ── COURSE INFO ────────────────────────────────────────────────
    add_modern_section_heading(doc, "Course Information")
    for k, v in course_info:
        p = doc.add_paragraph()
        rk = p.add_run(f"{k}:  ")
        rk.bold = True
        rk.font.size = Pt(11)
        rk.font.color.rgb = BLACK
        rv = p.add_run(v)
        rv.font.size = Pt(11)
        rv.font.color.rgb = BLACK
        p.paragraph_format.space_after = Pt(2)

    # ── TEAM ───────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Submitted By")
    team_table = doc.add_table(rows=len(team), cols=2)
    team_table.alignment = 0
    team_table.columns[0].width = Cm(4)
    team_table.columns[1].width = Cm(11)
    for i, (rid, name) in enumerate(team):
        ic, nc = team_table.rows[i].cells
        ic.text = ""
        nc.text = ""
        r1 = ic.paragraphs[0].add_run(rid)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = TEAL
        r2 = nc.paragraphs[0].add_run(name)
        r2.font.size = Pt(11)
        r2.font.color.rgb = BLACK
    for row in team_table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                borders.append(b)
            tc_pr.append(borders)

    # ── Footer ribbon: dark band + 2 teal parallelograms ───────────
    for _ in range(2):
        doc.add_paragraph()

    ribbon = doc.add_table(rows=1, cols=1)
    ribbon.autofit = True
    rib_cell = ribbon.cell(0, 0)
    shade_cell(rib_cell, DARK_HEX)
    rib_cell.text = ""
    rp = rib_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    spacer_run = rp.add_run("    ")
    spacer_run.font.size = Pt(14)
    # right-side teal mark using slash characters
    mark = rp.add_run("//")
    mark.font.size = Pt(28)
    mark.bold = True
    mark.font.color.rgb = TEAL
    pad = rp.add_run(" ")
    pad.font.size = Pt(14)
    # Remove ribbon borders
    tc_pr = rib_cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tc_pr.append(borders)

    doc.add_page_break()


def add_modern_section_heading(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(13.5)
    run.font.color.rgb = BLACK
    # add extra spacing before each heading except the first
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run.font.color.rgb = BLACK
    # subtle character spacing via XML
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "30")  # 30 = ~3pt extra spacing
    rPr.append(spacing)


# ─── PROPOSAL DOCUMENT ───────────────────────────────────────────────
def build_proposal() -> Document:
    doc = setup_document()

    # Modern one-page cover (matches the reference template)
    add_modern_cover(
        doc,
        big_title_lines=("Project", "Proposal"),
        author_lines=[
            "Aaliyan Asim, Rehan Kashif & Moeed Zahid",
            "April 2026",
        ],
        project_title="AI Study Assistant — A RAG-Powered Learning Companion",
        description=(
            "The AI Study Assistant is a smart web application that helps "
            "students study from their own PDFs (textbooks, lecture notes, "
            "slides) by answering questions, summarising content, generating "
            "multiple-choice quizzes, and giving simple explanations. It uses "
            "Retrieval-Augmented Generation (RAG) so that every answer is "
            "grounded in the uploaded material — not made up by the AI. "
            "Students upload their notes and chat with them; the app retrieves "
            "the most relevant pieces and asks a Large Language Model to "
            "answer using only those pieces."
        ),
        objective=(
            "To deliver a free, polished web app — usable by anyone via a "
            "single link — that lets students upload PDFs and interact with "
            "their study material through four scholarly modes: Inquire (Q&A), "
            "Summarise, Quiz Me (MCQs), and Explain Simply. The system must "
            "remain accurate, refuse to hallucinate, and run on free hosting."
        ),
        opportunity=(
            "Generic AI chatbots cannot read a student's actual notes and "
            "may invent answers. By combining open-weight LLaMA-3.1-8B "
            "(via free hosting on Groq) with a private vector index built "
            "from the user's own PDFs, we deliver personalised, trustworthy "
            "study help. The same architecture extends naturally to corporate "
            "training, customer support, and research assistants — making "
            "this a strong portfolio project with real business potential."
        ),
        course_info=[
            ("Course", "Natural Language Processing (NLP)"),
            ("Instructor", "Sir Adnan Abid"),
            ("Institution", "Punjab University College of Information Technology (PUCIT)"),
            ("Department", "Data Science"),
        ],
        team=[
            ("BSDSF23M042", "Aaliyan Asim"),
            ("BSDSF23M047", "Rehan Kashif"),
            ("BSDSF23M049", "Moeed Zahid"),
        ],
    )

    # ─── 1. Project Overview ─────────────────────────────────────────
    add_modern_section_heading(doc, "1. Project Overview")
    add_para(
        doc,
        "The AI Study Assistant is a smart web application that helps students "
        "study from their own PDFs (textbooks, lecture notes, slides) by "
        "answering questions, summarising content, generating multiple-choice "
        "quizzes, and giving simple explanations. It uses a technique called "
        "Retrieval-Augmented Generation (RAG) so that every answer is grounded "
        "in the uploaded material — not made up by the AI."
    )
    add_para(
        doc,
        "In short: students upload their notes, then chat with them. The app "
        "retrieves the most relevant pieces from those notes and asks a Large "
        "Language Model to write the answer using only those pieces.",
        italic=True, color=SOFT,
    )

    # ─── 2. Problem Statement ────────────────────────────────────────
    add_modern_section_heading(doc, "2. Problem Statement")
    add_para(
        doc,
        "Students often face common study challenges:"
    )
    for bp in [
        "Reading 100+ pages of slides before an exam is exhausting and inefficient.",
        "Generic AI chatbots like ChatGPT do not know your professor's specific notes — they may hallucinate.",
        "Making your own MCQs to practise is time-consuming.",
        "Some topics are explained too technically; learners need an ELI5 version.",
    ]:
        add_bullet(doc, bp)
    add_para(
        doc,
        "Our project solves these by combining the power of a Large Language "
        "Model with the student's actual material. The AI is forced to read "
        "the notes first, then answer — making it accurate, personal, and "
        "trustworthy."
    )

    # ─── 3. What the App Will Do ────────────────────────────────────
    add_modern_section_heading(doc, "3. What the App Will Do")
    add_para(doc, "The app provides four scholarly modes:", bold=True)
    add_kv_table(doc, [
        ("❓ Inquire (Q&A)",
         "Ask any question about your uploaded PDFs and get a structured answer with key points and the source it came from."),
        ("📝 Summarise",
         "Get a concise summary with the main themes, key concepts, and important details — perfect for last-minute revision."),
        ("📋 Quiz Me (MCQs)",
         "Auto-generate 5 multiple-choice questions with 4 options each and an answer key. Great for self-testing."),
        ("💡 Explain Simply",
         "An ELI5 explanation using everyday language and real-world analogies — useful for difficult topics."),
    ], headers=("Mode", "What it does"))

    add_callout(
        doc,
        "Anti-Hallucination Guard",
        "If a question is not covered in the uploaded material, the assistant "
        "will reply: \"The answer is not available in the provided material.\" "
        "It will never invent facts."
    )

    # ─── 4. Tools & Technologies ─────────────────────────────────────
    add_modern_section_heading(doc, "4. Tools and Technologies Used")
    add_para(doc, "Each piece of the system is mapped to a specific tool:")
    add_kv_table(doc, [
        ("Programming Language", "Python 3.12 — the most common language for AI/ML."),
        ("Web Framework (UI)", "Streamlit — turns Python code into a polished web app with no HTML/CSS/JS."),
        ("LLM (cloud)", "Groq running LLaMA-3.1-8B-Instant — fast and free for our usage."),
        ("LLM (local fallback)", "Ollama running LLaMA-3.1-8B — used when running offline on a developer's PC."),
        ("Embeddings (cloud)", "fastembed with BGE-small (BAAI) — converts text to vectors in-process."),
        ("Embeddings (local)", "Ollama nomic-embed-text — used in offline mode."),
        ("Vector Search", "Pure-Python NumPy cosine similarity — lightweight, no native libraries needed."),
        ("PDF Parsing", "pypdf — extracts text from each page of uploaded PDFs."),
        ("Version Control", "Git + GitHub (mmoeedz/ai-study-assistant)."),
        ("Hosting", "Streamlit Community Cloud — free public deployment so anyone can use the app via a link."),
    ], headers=("Component", "Tool / Library"))

    # ─── 5. How It Works (Architecture) ──────────────────────────────
    add_modern_section_heading(doc, "5. How It Works (in simple words)")
    add_para(doc, "The app has two main flows:", bold=True)

    add_para(doc, "A. Ingestion (when you upload a PDF):", bold=True)
    add_numbered(doc, "Read the PDF page by page using pypdf.")
    add_numbered(doc, "Cut the text into 1000-character chunks (with 200-char overlap).")
    add_numbered(doc, "Convert each chunk into a numerical vector (embedding).")
    add_numbered(doc, "Save the chunks + vectors into a small searchable index on disk.")

    add_para(doc, "B. Question answering (when you ask something):", bold=True)
    add_numbered(doc, "Convert your question into a vector.")
    add_numbered(doc, "Find the top-4 most similar chunks in the index using cosine similarity.")
    add_numbered(doc, "Build a prompt: instructions + retrieved chunks + your question.")
    add_numbered(doc, "Send to the LLM. It writes the answer using ONLY those chunks.")
    add_numbered(doc, "Show the answer with the source page numbers.")

    # ─── 6. Business Perspective ─────────────────────────────────────
    add_modern_section_heading(doc, "6. Business Perspective")
    add_para(doc, "Even though this is a course project, it has real-world value:")
    add_bullet(doc, "Target audience: university students, school students, MOOC learners, and self-taught professionals studying technical PDFs.")
    add_bullet(doc, "Pain it removes: cuts study time by giving instant explanations, summaries, and practice quizzes from a student's own material.")
    add_bullet(doc, "Privacy advantage: in local-mode, all data stays on the user's PC — important for confidential corporate training material.")
    add_bullet(doc, "Cost advantage: uses free models (Groq free tier, open-weights LLaMA, BGE embeddings) — operating cost is near zero.")
    add_bullet(doc, "Monetisation paths (future): freemium SaaS (limit free PDFs, charge for unlimited), B2B licence to coaching academies, or institutional licence for universities.")
    add_bullet(doc, "Scalability: thousands of students can use the deployed cloud version simultaneously without changing the architecture.")

    add_callout(
        doc,
        "Key Differentiators vs. Generic Chatbots",
        "1) Answers are grounded in YOUR documents (not the open internet).  "
        "2) Built-in MCQ generator and ELI5 mode — purpose-built for studying.  "
        "3) Open-source and free; users can self-host for full data privacy."
    )

    # ─── 7. Project Outcomes ─────────────────────────────────────────
    add_modern_section_heading(doc, "7. Expected Outcomes")
    for bp in [
        "A live, public web app accessible from any device via a single link.",
        "A reusable Python codebase published on GitHub.",
        "Hands-on experience with the full RAG pipeline: chunking → embeddings → vector search → prompting.",
        "An automated test suite proving accuracy and the no-hallucination rule.",
        "A polished, animated UI demonstrating attention to user experience.",
    ]:
        add_bullet(doc, bp)

    # ─── 8. Timeline ─────────────────────────────────────────────────
    add_modern_section_heading(doc, "8. Project Timeline")
    add_kv_table(doc, [
        ("Day 1", "Research RAG architecture; set up Python + Ollama environment."),
        ("Day 2", "Build PDF extraction + chunking + embeddings pipeline."),
        ("Day 3", "Build the vector store and retrieval logic."),
        ("Day 4", "Write prompt templates for the four modes and integrate the LLM."),
        ("Day 5", "Build the Streamlit UI with theme + animations."),
        ("Day 6", "Add cloud-mode (Groq) provider switch and run end-to-end tests."),
        ("Day 7", "Push to GitHub, deploy to Streamlit Cloud, and write documentation."),
    ], headers=("Phase", "Work Done"))

    # ─── 9. Conclusion ───────────────────────────────────────────────
    add_modern_section_heading(doc, "9. Conclusion")
    add_para(
        doc,
        "The AI Study Assistant proves that a small team (or even a single "
        "student) can use modern open-source AI tools to build a useful, "
        "polished product. By combining Streamlit's quick UI development, "
        "Groq's free LLM hosting, and the RAG technique, we deliver an app "
        "that genuinely helps students learn — while teaching us, the "
        "developers, the practical pipeline behind every modern AI assistant."
    )

    add_separator(doc)
    add_para(
        doc,
        "GitHub: https://github.com/mmoeedz/ai-study-assistant   |   "
        "Hosted on: Streamlit Community Cloud",
        size=10, italic=True, color=SOFT,
    )
    return doc


# ─── GUIDE DOCUMENT ──────────────────────────────────────────────────
def _add_guide_header(doc):
    """Modern header at the top of the guide (no full cover page)."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(11)
    table.columns[1].width = Cm(6)
    left, right = table.rows[0].cells

    # left: stacked PROJECT / GUIDE
    left.text = ""
    for line in ("Project", "Guide"):
        p = left.add_paragraph()
        run = p.add_run(line.upper())
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(34)
        run.font.color.rgb = BLACK
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 0.95

    # right meta block
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right.text = ""
    for i, line in enumerate([
        "AI Study Assistant",
        "Aaliyan Asim, Rehan Kashif & Moeed Zahid",
        "NLP — PUCIT — Sir Adnan Abid",
        "April 2026",
    ]):
        p = right.add_paragraph() if i > 0 else right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
        run.bold = (i == 0)
        p.paragraph_format.space_after = Pt(2)

    # remove all borders
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                borders.append(b)
            tc_pr.append(borders)

    rule_p = doc.add_paragraph()
    _set_paragraph_border(rule_p, color=TEAL_HEX, size_pt=18)
    rule_p.paragraph_format.space_before = Pt(2)
    rule_p.paragraph_format.space_after = Pt(14)


def _add_code_block(doc, code: str):
    """Render a fixed-width code block with a light-gray background."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, "F2F4F7")
    set_cell_borders(cell, "D0D5DD")
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def build_guide() -> Document:
    doc = setup_document()

    # No cover page — modern header + teal rule, then straight into content.
    _add_guide_header(doc)

    add_modern_section_heading(doc, "Welcome")
    add_para(
        doc,
        "This guide walks through every part of the AI Study Assistant "
        "project in plain language. By the end, you will understand what "
        "each file does, why we picked the tools we did, and exactly how all "
        "the pieces fit together. We have written it like a friend explaining "
        "the project — no jargon walls, no skipped steps, no hand-waving. "
        "Wherever a piece of code matters, we show it. Wherever a concept is "
        "tricky, we use an analogy."
    )

    # ───────────────────────────────────────────────────────────────────
    # Step 1 — The Idea
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 1 — The Idea")
    add_para(
        doc,
        "Most students struggle with the same problem: too much reading, "
        "too little time. We get a stack of PDFs (slides, papers, books) "
        "and exam day creeps closer. Generic chatbots like ChatGPT are not "
        "much help here because they have not read OUR specific notes — "
        "and worse, they sometimes make up confident-sounding answers."
    )
    add_para(
        doc,
        "So we asked ourselves: what if the AI was forced to read OUR own "
        "notes first, and only answer using those notes? That is the core "
        "idea behind the AI Study Assistant. Once that 'private knowledge' "
        "was working, we wrapped four useful study modes around it:"
    )
    add_bullet(doc, "INQUIRE — ask any question; get a structured answer with citations.")
    add_bullet(doc, "SUMMARISE — get the gist of a chapter or whole document.")
    add_bullet(doc, "QUIZ ME — auto-generate 5 MCQs with an answer key, perfect for self-testing.")
    add_bullet(doc, "EXPLAIN SIMPLY — ELI5 explanations using analogies and simple language.")
    add_para(
        doc,
        "Each mode is just a different prompt template that we send to the "
        "AI along with the relevant pieces of YOUR PDFs. The 'brain' is "
        "shared — only the 'instructions' change."
    )

    # ───────────────────────────────────────────────────────────────────
    # Step 2 — Picking the Tools
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 2 — Picking the Tools")
    add_para(
        doc,
        "Before writing a single line of code we listed every capability "
        "the project would need, then chose the simplest free tool for each:"
    )
    add_kv_table(doc, [
        ("A web UI",
         "Streamlit — write plain Python, get a polished web app. Zero HTML/CSS/JS required."),
        ("A brain (LLM)",
         "LLaMA-3.1-8B-Instant. Free via Groq (cloud) or Ollama (local)."),
        ("Read PDFs",
         "pypdf — pure-Python PDF reader. No native libraries, works on Windows, macOS, Linux."),
        ("Find relevant chunks",
         "fastembed (BGE-small) for cloud OR Ollama nomic-embed-text for local. Both create 'embeddings' (number lists that capture meaning)."),
        ("Vector search",
         "Tiny custom NumPy index (cosine similarity). Avoids heavy native libraries like FAISS that some Windows machines block."),
        ("Hosting",
         "Streamlit Community Cloud — free public link, redeploys automatically when GitHub updates."),
        ("Source control",
         "Git + GitHub. Public repo so anyone can read, fork, or contribute."),
    ], headers=("Need", "Our Choice"))

    add_callout(
        doc,
        "Why two LLM providers?",
        "Cloud mode (Groq) lets ANYONE use the app via a single link — perfect "
        "for sharing with friends or showing in a portfolio. Local mode (Ollama) "
        "keeps everything on the user's PC — no internet required, no API key, "
        "complete privacy. Both modes share the same code; only one setting "
        "changes (LLM_PROVIDER)."
    )

    # ───────────────────────────────────────────────────────────────────
    # Step 3 — The Project Files
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 3 — The Project Files")
    add_para(
        doc,
        "We deliberately kept the file count low. Here is every important "
        "file in the repository and exactly what lives inside it:"
    )
    add_kv_table(doc, [
        ("app.py",
         "The Streamlit web app — UI layout, custom CSS theme, sidebar with file upload, mode selector, chat input, and response cards. ~700 lines of Python (most of which is CSS in a string)."),
        ("rag_pipeline.py",
         "The 'brain' of the project. Contains: Document data class, RecursiveCharacterTextSplitter (chunker), SimpleVectorStore (NumPy-based index), OllamaClient and GroqClient (LLM HTTP), FastEmbedEmbedder, and the StudyAssistant class that ties everything together."),
        ("prompts.py",
         "Four prompt templates (one per mode). Each contains strict 'no hallucination' rules + the exact output format we want."),
        ("config.py",
         "Central settings: provider switch, model names, chunk size, retrieval top-k, temperature. Reads Streamlit secrets and environment variables."),
        ("requirements.txt",
         "Python dependencies: streamlit, numpy, pypdf, fastembed, requests."),
        (".streamlit/config.toml",
         "Streamlit theme (navy/gold colours) + server defaults."),
        (".streamlit/secrets.toml",
         "API keys (e.g. GROQ_API_KEY). GIT-IGNORED — never pushed to GitHub."),
        (".streamlit/secrets.toml.example",
         "Template showing what keys are expected."),
        ("tests/test_e2e.py",
         "Eight automated tests proving every mode works and the AI refuses to hallucinate."),
        ("tests/smoke_groq.py",
         "Quick test of the cloud-mode pipeline (no PDFs needed)."),
        ("vectorstore/",
         "Auto-created folder where embeddings + chunks are pickled to disk."),
        ("README.md",
         "GitHub project page with overview, setup instructions, and architecture."),
        (".gitignore",
         "Excludes: __pycache__, .venv, secrets.toml, vectorstore, fastembed_cache, log files."),
    ], headers=("File / Folder", "What it contains"))

    # ───────────────────────────────────────────────────────────────────
    # Step 4 — How RAG Works
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 4 — How RAG Works")
    add_para(
        doc,
        "RAG stands for Retrieval-Augmented Generation. It is the technique "
        "that makes the assistant answer based on YOUR documents instead of "
        "from the LLM's general training data. It is just three steps:"
    )
    add_numbered(doc, "RETRIEVAL — search the user's PDFs for the parts most related to the question.")
    add_numbered(doc, "AUGMENTATION — paste those relevant parts into the prompt as 'context'.")
    add_numbered(doc, "GENERATION — ask the LLM to write the answer using ONLY that context.")

    add_para(doc, "Worked example:", bold=True)
    add_para(
        doc,
        "User asks: 'What is photosynthesis?'"
    )
    add_numbered(doc,
        "We turn the question into an embedding (a list of 384 numbers).")
    add_numbered(doc,
        "We compare that embedding to every chunk's embedding in the user's PDF index, and pick the top 4 most similar chunks. One of them might be: 'Plants convert sunlight to chemical energy through photosynthesis, occurring mainly in the chloroplasts...'")
    add_numbered(doc,
        "We build a prompt: \"Use only this context. CONTEXT: [chunk 1, chunk 2, ...]. QUESTION: What is photosynthesis?\"")
    add_numbered(doc,
        "The LLM generates a structured response (Answer / Key Points / Source Insight) using only those 4 chunks. It also cites the source page numbers.")

    add_callout(
        doc,
        "Why is RAG better than just asking the LLM?",
        "1) The answer is grounded in YOUR specific notes, not the open web. "
        "2) Hallucinations drop drastically because the LLM is given the facts. "
        "3) Source citations are automatic — you can verify every answer. "
        "4) New information (e.g. last week's lecture) works instantly without "
        "retraining the model."
    )

    # ───────────────────────────────────────────────────────────────────
    # Step 5 — Chunks and Embeddings
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 5 — Chunks, Embeddings & Cosine Similarity")

    add_para(doc, "What is a chunk?", bold=True)
    add_para(
        doc,
        "An LLM has a limited 'context window' — it cannot read a 100-page "
        "PDF in one go. So we cut the text into roughly 1000-character "
        "pieces called CHUNKS. Each chunk overlaps the next one by 200 "
        "characters so a sentence cut in half does not lose its meaning."
    )
    add_para(doc, "Settings used in this project:", italic=True, color=GRAY)
    add_bullet(doc, "CHUNK_SIZE = 1000 characters (~150–200 words).")
    add_bullet(doc, "CHUNK_OVERLAP = 200 characters.")
    add_bullet(doc, "Splitter algorithm: try paragraph break first, then newline, then full-stop, then space — only as a last resort cut by character count.")

    add_para(doc, "Why exactly these numbers?", bold=True)
    add_para(
        doc,
        "Chunk size and overlap are the two most-tuned knobs in any RAG "
        "system. Here is the reasoning behind ours:"
    )
    add_bullet(doc,
        "1000 chars (~150–200 words) is the 'sweet spot' that academic RAG "
        "papers and tools like LangChain default to. SMALLER (200 chars): "
        "loses context — a definition can get split from its example. "
        "LARGER (3000 chars): the LLM gets confused by too much irrelevant "
        "text, and we waste tokens on the API.")
    add_bullet(doc,
        "200-char overlap (~20% of chunk size) is enough to keep the last "
        "sentence of one chunk visible at the start of the next. Without "
        "overlap, a sentence cut in half loses its meaning. With too much "
        "overlap, we duplicate content and waste storage.")
    add_bullet(doc,
        "Splitting on paragraph → newline → full-stop → space (in that "
        "order) is from LangChain's RecursiveCharacterTextSplitter. It tries "
        "to break at NATURAL boundaries first, so we never cut mid-sentence "
        "unless we absolutely have to.")
    add_bullet(doc,
        "We retrieve TOP-4 chunks (k=4) per query. Lower (k=1): risk of "
        "missing the relevant passage. Higher (k=10): the LLM gets "
        "distracted, and the prompt becomes huge. k=4 gives ~4000 chars of "
        "context — comfortable for an 8K-token model.")

    add_para(doc, "What is an embedding?", bold=True)
    add_para(
        doc,
        "An embedding is a list of numbers that represents the MEANING of "
        "a piece of text. The model used in our cloud mode (BAAI/bge-small) "
        "produces 384 numbers per chunk; the local model (nomic-embed-text) "
        "produces 768. The exact numbers do not matter on their own — what "
        "matters is that two chunks with similar meaning have similar number "
        "lists. This lets us search by meaning, not by exact words."
    )
    add_para(doc, "Mini example:", italic=True, color=GRAY)
    add_bullet(doc, "'cars are vehicles' → [0.12, -0.04, 0.81, ..., 0.07]")
    add_bullet(doc, "'automobiles transport people' → similar numbers (close in meaning)")
    add_bullet(doc, "'pancakes for breakfast' → very different numbers (no semantic overlap)")

    add_para(doc, "Why these specific embedding models?", bold=True)
    add_bullet(doc,
        "BAAI/bge-small-en-v1.5 (cloud mode) — only 33 MB, runs in-process via "
        "fastembed (ONNX runtime). It scores in the top-10 on the MTEB English "
        "leaderboard among models <100 MB, so we get high quality without "
        "needing a GPU or paid API.")
    add_bullet(doc,
        "nomic-embed-text (local Ollama mode) — already pulled by Ollama users "
        "for other RAG demos, supports long inputs, and has a 768-dim output. "
        "We keep it for users who already have Ollama running.")
    add_bullet(doc,
        "We did NOT use OpenAI text-embedding-ada — it requires a paid API key "
        "and would tie us to one vendor.")

    add_para(doc, "What is cosine similarity?", bold=True)
    add_para(
        doc,
        "Imagine two arrows starting from the same point. If they point in "
        "the same direction, the cosine of the angle between them is 1 (very "
        "similar). If they are perpendicular, it is 0 (unrelated). If they "
        "point opposite, it is -1 (opposite meaning). Cosine similarity is "
        "the maths formula that gives us this number for any two embeddings."
    )
    add_para(doc, "Used in our retrieval step:", italic=True, color=GRAY)
    _add_code_block(doc,
        "similarities = embeddings @ query_vector / (||embeddings|| * ||query||)\n"
        "top_k_indices = argsort(similarities)[-4:][::-1]\n"
        "return [documents[i] for i in top_k_indices]"
    )

    # ───────────────────────────────────────────────────────────────────
    # Step 6 — Prompt Templates
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 6 — Prompt Engineering")
    add_para(
        doc,
        "The single most important file for output quality is prompts.py. "
        "Each of the four modes has its own carefully written template. "
        "Every template has the same four parts:"
    )
    add_numbered(doc, "STRICT RULES — 'Use only the provided context. Do not invent. If the answer is not in the context, say so exactly.'")
    add_numbered(doc, "CONTEXT — the top-4 chunks we retrieved, with their source filename and page number.")
    add_numbered(doc, "QUESTION — what the user asked.")
    add_numbered(doc, "OUTPUT FORMAT — the exact section structure (e.g. ✅ Answer / 📌 Key Points / 📖 Source Insight).")

    add_para(doc, "Example: the QA template (simplified)", italic=True, color=GRAY)
    _add_code_block(doc,
        'You are an AI Study Assistant. You MUST follow these rules:\n'
        '1. Use ONLY the context below. Do NOT use outside knowledge.\n'
        '2. If the answer is not in the context, reply EXACTLY:\n'
        '   "The answer is not available in the provided material."\n'
        '\n'
        'Context:\n'
        '{context}\n'
        '\n'
        'Question:\n'
        '{question}\n'
        '\n'
        'Format your response in this exact structure:\n'
        '### ✅ Answer:\n'
        '[detailed answer]\n'
        '### 📌 Key Points:\n'
        '• [point 1]\n'
        '• [point 2]\n'
        '### 📖 Source Insight:\n'
        '[which chunks/pages were used]'
    )
    add_para(
        doc,
        "The MCQ template asks for exactly 5 questions with options A–D and "
        "an Answer Key block. The ELI5 template asks for analogies and "
        "everyday language. The Summarise template asks for themes and key "
        "concepts. Same retrieval pipeline — just a different blueprint."
    )

    # ───────────────────────────────────────────────────────────────────
    # Step 7 — UI
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 7 — Building the Web Interface")
    add_para(
        doc,
        "The UI lives entirely in app.py. It is a single Python file but "
        "looks polished thanks to ~400 lines of custom CSS injected through "
        "st.markdown(...). The layout has three regions:"
    )
    add_bullet(doc, "HEADER — animated gold crest, large 'AI Study Assistant' title, italic subtitle, gold horizontal rule with a centred ✦.")
    add_bullet(doc, "SIDEBAR ('The Library') — file upload, 'Process Documents' button with progress bar, two stat plates (Documents / Chunks), list of indexed PDFs, Clear All Data button, footer credit.")
    add_bullet(doc, "MAIN AREA — bookmark-style mode selector, chat history, parchment-style response cards with gold-leaf left edge, source-citation chips, and an expander showing retrieved chunks.")

    add_para(doc, "Theme details:", bold=True)
    add_bullet(doc, "Palette: deep navy background (#0a1322), antique gold accents (#d4a84b), parchment cards (#f5ecd7).")
    add_bullet(doc, "Typography: Cormorant Garamond serif for headings (academic feel), Inter sans-serif for body.")
    add_bullet(doc, "Animations: glow on the crest, page-open rotateX entrance on the welcome card, slide-in on chat bubbles, hover lift on stat plates and source chips.")
    add_bullet(doc, "Hidden by CSS: Streamlit's default top-bar, footer, and 'Made with Streamlit' watermark.")

    add_para(doc, "How session state works:", bold=True)
    add_para(
        doc,
        "Streamlit re-runs the entire app.py on every interaction. To keep "
        "data alive between re-runs we use st.session_state to hold: the "
        "StudyAssistant instance, the chat history, and the 'processed' "
        "flag. The vector index itself lives on disk (in vectorstore/) so "
        "it survives even when the server restarts."
    )

    # ───────────────────────────────────────────────────────────────────
    # Step 8 — Testing
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 8 — Verification & Testing")
    add_para(
        doc,
        "We wrote an end-to-end test suite that proves every feature works. "
        "It uses real PDFs from the Slides folder, runs ingestion through "
        "the actual Ollama models, and checks the output of every mode. "
        "All eight tests pass. Run with: python -X utf8 tests/test_e2e.py."
    )
    add_kv_table(doc, [
        ("Clean state",
         "PASS — vector store can be reset; total_chunks back to 0."),
        ("Ingest",
         "PASS — 2 PDFs split into 133 chunks in ~270s, with progress callbacks fired."),
        ("Q&A mode",
         "PASS — response contains all three required markers (✅ Answer / 📌 Key Points / 📖 Source Insight); 4 sources returned."),
        ("Summarise mode",
         "PASS — response contains 📝 Summary, 🎯 Main Themes, 📌 Key Concepts sections."),
        ("Quiz Me mode",
         "PASS — response contains 5 questions with options A–D and an Answer Key block."),
        ("Explain Simply mode",
         "PASS — response contains 🧒 Simple Explanation, 🔑 Remember These, 💡 Real-World Analogy."),
        ("Hallucination guard",
         "PASS — when asked an off-topic chemistry question, the model refused with 'I cannot provide information…' rather than making up an answer."),
        ("Persistence",
         "PASS — after clearing the in-memory assistant and creating a new one, the saved index reloaded from disk and retrieval still worked."),
    ], headers=("Test", "Result"))

    # ───────────────────────────────────────────────────────────────────
    # Step 9 — Deployment
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 9 — Deployment Story")
    add_para(
        doc,
        "Getting the app to a public URL was the trickiest part because of "
        "platform constraints:"
    )
    add_bullet(doc, "Vercel cannot run Streamlit (Vercel only supports short-lived serverless functions; Streamlit needs a long-running Python process).")
    add_bullet(doc, "Streamlit Cloud cannot run Ollama (no local LLM, can't ship a 5-GB model file in a free tier container).")
    add_bullet(doc, "FAISS native libraries are blocked by some corporate Application Control policies on Windows.")
    add_para(doc, "Our solutions:", bold=True)
    add_numbered(doc, "Made the LLM provider swappable. The same StudyAssistant class can call Ollama (local mode) or Groq (cloud mode) depending on a single setting.")
    add_numbered(doc, "Replaced Ollama embeddings with fastembed (a tiny ONNX model that runs in-process). This means the cloud version needs zero local services.")
    add_numbered(doc, "Replaced FAISS with a pure-Python NumPy cosine-similarity index. Avoids native dependencies; performance is still great for thousands of chunks.")
    add_numbered(doc, "Pushed the code to GitHub: github.com/mmoeedz/ai-study-assistant.")
    add_numbered(doc, "Connected Streamlit Cloud to the repo. It auto-builds on every push.")
    add_numbered(doc, "Added GROQ_API_KEY to Streamlit Cloud's Secrets dashboard (NOT to GitHub).")
    add_para(doc, "Result: a free, public URL — anyone can use the app from any device.", bold=True)

    add_callout(
        doc,
        "What about ngrok / Cloudflare Tunnel?",
        "Both tools can expose a locally running app to the public internet. "
        "We did NOT use them for the final deploy because they require the "
        "developer's PC to stay on. The Streamlit Cloud + Groq solution "
        "removes that dependency entirely."
    )

    # ───────────────────────────────────────────────────────────────────
    # Step 10 — Lessons
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 10 — Lessons Learnt")
    for bp in [
        "RAG is conceptually simple — three steps. The hard parts are good chunking and good prompts.",
        "Embeddings give us 'search by meaning'; cosine similarity ranks them.",
        "Streamlit lets a small team produce a polished AI demo astonishingly fast.",
        "Prompt engineering is THE main lever for output quality and safety.",
        "Designing for two providers (cloud + local) makes the project flexible and demo-friendly.",
        "Always gitignore secrets — never commit API keys to GitHub. We rotated our Groq key once after accidental exposure.",
        "Cloud deployment forces re-thinking heavy local services like Ollama; lighter alternatives exist for almost everything.",
        "A clean UI is not optional — it transforms how the project is perceived.",
        "Automated tests pay off: we caught regressions immediately when refactoring the embeddings backend.",
    ]:
        add_bullet(doc, bp)

    # ───────────────────────────────────────────────────────────────────
    # Step 11 — Design Decisions & Trade-offs (the WHY behind every choice)
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 11 — Design Decisions & Trade-offs")
    add_para(
        doc,
        "Every project is a chain of decisions. For each major choice we "
        "considered the alternatives, picked the simplest one that satisfied "
        "our requirements, and accepted the trade-off. This section explains "
        "the reasoning behind each one."
    )

    # — UI framework
    add_para(doc, "Why Streamlit (not Flask, FastAPI, or React)?", bold=True)
    add_bullet(doc, "Streamlit lets us write 100% Python — no HTML, CSS, or JavaScript skills required for the basic version.")
    add_bullet(doc, "Built-in widgets for file upload, chat, sidebar, progress bars — exactly what we needed.")
    add_bullet(doc, "Has a free official cloud host (Streamlit Community Cloud), making deployment trivial.")
    add_bullet(doc, "Trade-off: less customisable than React, but for an internal tool / demo this is a feature, not a bug.")

    # — LLM choice
    add_para(doc, "Why LLaMA-3.1-8B (not GPT-4 or LLaMA-70B)?", bold=True)
    add_bullet(doc, "LLaMA-3.1-8B is FREE on Groq (cloud) and small enough to run on a laptop via Ollama (local).")
    add_bullet(doc, "It is good enough for educational Q&A — modern 8B models are roughly equal to GPT-3.5 in quality.")
    add_bullet(doc, "GPT-4 / Claude would cost real money per request and tie us to a paid vendor — bad for a free student tool.")
    add_bullet(doc, "70B models would be too heavy for laptop hardware in local mode.")

    # — Provider strategy
    add_para(doc, "Why dual providers (Groq + Ollama)?", bold=True)
    add_bullet(doc, "Groq alone: works for everyone but requires internet + an API key.")
    add_bullet(doc, "Ollama alone: zero API costs and full privacy, but needs every user to install and pull 5 GB of models.")
    add_bullet(doc, "BOTH: cloud users get a public URL with no install; privacy-conscious users (or offline classrooms) can run it 100% locally.")
    add_bullet(doc, "Trade-off: extra abstraction in rag_pipeline.py, but worth it for the deployment flexibility.")

    # — Vector store
    add_para(doc, "Why a custom NumPy index (not FAISS or Chroma)?", bold=True)
    add_bullet(doc, "FAISS requires a native C++ library (faiss-cpu / faiss-gpu) — blocked on some Windows machines by Application Control policies.")
    add_bullet(doc, "Chroma adds a database server dependency (sqlite or duckdb) and writes thousands of files.")
    add_bullet(doc, "Our pure-Python NumPy approach is ~50 lines of code, has zero native dependencies, and is fast enough for thousands of chunks.")
    add_bullet(doc, "Trade-off: would slow down at millions of chunks, but no student PDF library reaches that scale.")

    # — Embeddings
    add_para(doc, "Why fastembed for cloud (not the Groq embedding API)?", bold=True)
    add_bullet(doc, "Groq does NOT currently offer an embeddings endpoint — only chat completions.")
    add_bullet(doc, "OpenAI / Cohere embeddings require a paid API key and add latency per chunk.")
    add_bullet(doc, "fastembed runs in-process on CPU. The 33 MB model downloads once and embeds 100 chunks/second on a laptop.")
    add_bullet(doc, "Trade-off: an extra ~70 MB on first deploy, but eliminates network calls during ingestion.")

    # — PDF parser
    add_para(doc, "Why pypdf (not PyMuPDF or pdfplumber)?", bold=True)
    add_bullet(doc, "pypdf is pure-Python — no native libraries, works identically on Windows / macOS / Linux.")
    add_bullet(doc, "PyMuPDF is faster and extracts tables better, but its AGPL license complicates distribution.")
    add_bullet(doc, "For text-only extraction from typical lecture slides and notes, pypdf is more than adequate.")

    # — Chunk size
    add_para(doc, "Why 1000 chars (and not 500 or 2000)?", bold=True)
    add_bullet(doc, "500 chars: chunks become too narrow — definitions get separated from their explanations, and retrieval misses context.")
    add_bullet(doc, "2000 chars: each chunk dominates the LLM's prompt. With top-4 we'd be sending 8000 chars of context — close to the model's limit.")
    add_bullet(doc, "1000 chars (≈150–200 words ≈ 1 paragraph) holds one COMPLETE idea, leaving room for 4 chunks + the prompt overhead.")
    add_bullet(doc, "This is also LangChain's default — battle-tested across thousands of community RAG projects.")

    # — Overlap
    add_para(doc, "Why 200-char overlap (and not 0 or 500)?", bold=True)
    add_bullet(doc, "0 overlap: a sentence split exactly at the chunk boundary loses its meaning entirely. The retrieval will never find it.")
    add_bullet(doc, "500 overlap: 50% of every chunk is duplicated. Wastes storage and inflates similarity scores artificially.")
    add_bullet(doc, "200 chars (≈20%) is the documented sweet spot — preserves cross-boundary sentences without major duplication.")

    # — Top-k
    add_para(doc, "Why top-k = 4 (not 1 or 10)?", bold=True)
    add_bullet(doc, "k=1: works for direct lookups but fails when the answer needs synthesis across multiple chunks.")
    add_bullet(doc, "k=10: too much context — the LLM has to read 10,000+ chars before answering. Slower, more expensive, and more likely to drift.")
    add_bullet(doc, "k=4: covers most multi-chunk questions while keeping the prompt small. Output quality plateaus around k=4–6 in our tests.")

    # — Temperature
    add_para(doc, "Why LLM temperature = 0.3?", bold=True)
    add_bullet(doc, "Temperature controls randomness. 0.0 = deterministic, 1.0 = creative writing.")
    add_bullet(doc, "0.3 keeps answers consistent and factual (re-asking the same question gives essentially the same answer), but allows tiny phrasing variations.")
    add_bullet(doc, "We tested 0.0 — answers were correct but stiff. Tested 0.7 — model started embellishing facts. 0.3 is the sweet spot.")

    # — Context window
    add_para(doc, "Why num_ctx = 4096 (not 8192)?", bold=True)
    add_bullet(doc, "4 chunks × 1000 chars + prompt instructions + question + response = roughly 3000 tokens. 4096 fits comfortably.")
    add_bullet(doc, "Larger context windows take more RAM in local mode (Ollama allocates the buffer up-front).")
    add_bullet(doc, "If a user uploads VERY long-context queries we can bump this — it's a single config change.")

    # — Hosting
    add_para(doc, "Why Streamlit Community Cloud (not Vercel, AWS, Heroku)?", bold=True)
    add_bullet(doc, "Vercel can't run Streamlit (Vercel only supports stateless serverless functions; Streamlit needs a long-running process).")
    add_bullet(doc, "AWS / GCP would require a paid VM and DevOps work — overkill for a student demo.")
    add_bullet(doc, "Heroku free tier was discontinued in 2022.")
    add_bullet(doc, "Streamlit Community Cloud auto-deploys on every git push, has free HTTPS, and is PURPOSE-BUILT for Streamlit apps.")

    # — Source control
    add_para(doc, "Why GitHub public (not private)?", bold=True)
    add_bullet(doc, "Streamlit Community Cloud's free tier requires a PUBLIC repo.")
    add_bullet(doc, "It also lets us share the project as a portfolio piece.")
    add_bullet(doc, "Sensitive files (secrets.toml, .env) are excluded via .gitignore — only the code is public, never the keys.")

    # — Anti-hallucination strategy
    add_para(doc, "Why explicit 'no hallucination' rules in the prompt?", bold=True)
    add_bullet(doc, "Without strict rules, LLMs WILL fall back to general knowledge when retrieval is weak — even when told to use only the context.")
    add_bullet(doc, "We pre-load the prompt with: 'ONLY use the context. If the answer is not present, reply EXACTLY: \"The answer is not available in the provided material.\"'")
    add_bullet(doc, "Our hallucination test (asking about sulfuric acid manufacture from an NLP PDF) confirmed the model correctly REFUSES instead of inventing.")

    add_callout(
        doc,
        "Summary heuristic",
        "When in doubt, we picked the option that: (a) was free, (b) had no native "
        "dependencies, (c) had broad community support, and (d) kept the user's "
        "data either local or under a single signed-in account. Almost every "
        "decision above falls out of those four constraints."
    )

    # ───────────────────────────────────────────────────────────────────
    # Step 12 — How to Run
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 12 — How to Run It Yourself")

    add_para(doc, "Cloud mode (no install — recommended for non-developers):", bold=True)
    add_para(doc,
        "Just open the deployed Streamlit URL in any browser. Upload your "
        "PDFs, click 'Process Documents', then start chatting.")

    add_para(doc, "Local mode with Ollama (no internet, full privacy):", bold=True)
    _add_code_block(doc,
        "git clone https://github.com/mmoeedz/ai-study-assistant\n"
        "cd ai-study-assistant\n"
        "pip install -r requirements.txt\n"
        "ollama pull llama3.1:8b\n"
        "ollama pull nomic-embed-text\n"
        "echo 'LLM_PROVIDER = \"ollama\"' > .streamlit/secrets.toml\n"
        "streamlit run app.py"
    )

    add_para(doc, "Local mode with Groq (cloud LLM but local UI):", bold=True)
    _add_code_block(doc,
        "git clone https://github.com/mmoeedz/ai-study-assistant\n"
        "cd ai-study-assistant\n"
        "pip install -r requirements.txt\n"
        "# Get a free key from https://console.groq.com/keys\n"
        "@'\n"
        'LLM_PROVIDER = "groq"\n'
        'GROQ_API_KEY = "gsk_your_key_here"\n'
        'GROQ_MODEL   = "llama-3.1-8b-instant"\n'
        "'@ | Set-Content .streamlit\\secrets.toml\n"
        "streamlit run app.py"
    )

    # ───────────────────────────────────────────────────────────────────
    # Step 13 — Final thoughts
    # ───────────────────────────────────────────────────────────────────
    add_modern_section_heading(doc, "Step 13 — Final Thoughts")
    add_para(
        doc,
        "This project shows what a small team can build end-to-end with "
        "modern open-source AI tools in about a week: a real, working, "
        "free-to-use product that solves a genuine student problem and "
        "looks professional on a CV. We hope this guide makes every "
        "internal detail clear and gives you the confidence to build your "
        "own RAG-powered apps. Happy studying — and happy hacking.",
    )
    add_para(doc, "Repository: github.com/mmoeedz/ai-study-assistant", italic=True, color=GRAY)
    add_para(doc, "Live demo: hosted on Streamlit Community Cloud", italic=True, color=GRAY)

    return doc


def export_pdf(docx_path: Path) -> Path | None:
    """Convert a .docx to .pdf using docx2pdf (requires MS Word on Windows)."""
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert  # type: ignore
        convert(str(docx_path), str(pdf_path))
        return pdf_path
    except Exception as e:
        print(f"⚠️  PDF export failed for {docx_path.name}: {e}")
        return None


# ─── Main ────────────────────────────────────────────────────────────
if __name__ == "__main__":
    out_dir = PROJECT_ROOT
    proposal_path = out_dir / "AI_Study_Assistant_Proposal.docx"
    guide_path = out_dir / "AI_Study_Assistant_Guide.docx"

    proposal = build_proposal()
    proposal.save(str(proposal_path))
    print(f"✅ Saved proposal -> {proposal_path}")

    guide = build_guide()
    guide.save(str(guide_path))
    print(f"✅ Saved guide    -> {guide_path}")

    # PDF export (guide only — that's what the user asked for)
    pdf = export_pdf(guide_path)
    if pdf:
        print(f"✅ Saved PDF      -> {pdf}")
